import docx

#Escribir
doc = docx.Document()
doc.add_paragraph('Hola GAturroXd! que mas te digo hoy en dia?'
                  )
doc.save('doc.docx')

# Leer
document = docx.Document("doc.docx")

# Recorre los párrafos del documento e imprime su contenido
for paragraph in document.paragraphs:
    print(paragraph.text)
